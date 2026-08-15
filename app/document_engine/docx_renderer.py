"""
Генерация настоящего DOCX-файла из content_blocks.

Требование спецификации (п.15): DOCX должен быть полноценной
редактируемой структурой (paragraphs/headings/alignment/spacing), а не
картинкой внутри Word. Используем python-docx напрямую — открывается
в Word/Google Docs/LibreOffice.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm


def render_docx(title: str, content_blocks: list[dict]) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title_p = doc.add_heading(title, level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for block in content_blocks:
        block_type = block.get("type")
        text = block.get("text", "")

        if block_type == "spacer":
            doc.add_paragraph()
            continue

        if block_type == "heading":
            h = doc.add_heading(text, level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if block_type == "heading_center":
            h = doc.add_heading(text, level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if block_type == "paragraph_right":
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            continue

        if block_type == "signature_line":
            p = doc.add_paragraph()
            p.add_run(f"{text}    ").bold = False
            p.add_run("_______________ /_______________/")
            continue

        # paragraph (по умолчанию)
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
