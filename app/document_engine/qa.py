"""
Техническая QA-проверка сгенерированных DOCX/PDF (п.41 спецификации).

Это НЕ AI-анализ содержания — только структурная проверка: файл
реально открывается, есть текст, нет очевидных признаков поломки.
Смысловая проверка (противоречия, орфография) требует AI и сюда не
входит — честно помечена как недоступная без ключа (см. app/ai/provider.py).
"""
from __future__ import annotations

import io

from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentQAError(Exception):
    pass


def check_docx(data: bytes) -> dict:
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise DocumentQAError(f"DOCX не открывается: {exc}") from exc

    paragraph_count = len(doc.paragraphs)
    has_text = any(p.text.strip() for p in doc.paragraphs)
    if not has_text:
        raise DocumentQAError("DOCX сгенерирован без текста")

    return {"opens": True, "paragraphs": paragraph_count, "has_text": has_text}


def check_pdf(data: bytes) -> dict:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise DocumentQAError(f"PDF не открывается: {exc}") from exc

    page_count = len(reader.pages)
    if page_count == 0:
        raise DocumentQAError("PDF сгенерирован без страниц")

    has_text = False
    empty_pages = 0
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            has_text = True
        else:
            empty_pages += 1

    if not has_text:
        raise DocumentQAError("PDF не содержит извлекаемого текста")

    return {"opens": True, "pages": page_count, "has_text": has_text, "empty_pages": empty_pages}
